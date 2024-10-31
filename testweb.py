import requests
from bs4 import BeautifulSoup
from transformers import pipeline
from docx import Document

# Define a dictionary of genres and corresponding URLs
genres_urls = {
    "ballet": "https://www.thelewisfoundation.org/2023/04/10-tips-for-beginners-ballet/",
    "gentle reminders": "https://vocal.media/motivation/10-reminders-you-need-to-hear-today-if-you-are-suffering-from-a-burnout",
    "journal prompts": "https://psychcentral.com/blog/ready-set-journal-64-journaling-prompts-for-self-discovery#the-journal-prompts",
    "productivity": "https://www.strangecharmed.com/2016/1-secret-no-one-told-productivity/",
}

# Function to display available genres and get user selection
def get_genre_choice():
    print("Available genres:")
    for genre in genres_urls.keys():
        print(f"- {genre}")
    genre_choice = input("\nEnter a genre from the list above: ").strip().lower()
    if genre_choice in genres_urls:
        return genres_urls[genre_choice]
    else:
        print("Invalid genre selected.")
        return None

# Function to extract body text from the URL
def extract_body_text(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
    }
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        body_text = soup.body.get_text(separator=' ', strip=True)
        return body_text
    else:
        return "Failed to retrieve content"

# Function to write the summary in a formatted docx file
def write_summary_to_doc(summary_text, genre):
    document = Document()
    document.add_heading(f"Summary of {genre.capitalize()}", level=1)
    document.add_paragraph(summary_text)
    file_name = f"{genre}_summary.docx"
    document.save(file_name)
    print(f"\nSummary saved to {file_name}")

# Get the URL based on user's genre choice
url = get_genre_choice()
if url:
    # Get the body text of the selected URL
    text = extract_body_text(url)

    # Load a summarization pipeline
    summarizer = pipeline("summarization", truncation=True)

    # Generate summary
    summary = summarizer(text, max_length=400, min_length=250, do_sample=False)
    summary_text = summary[0]['summary_text']

    # Write the summary to a formatted .docx file
    genre = list(genres_urls.keys())[list(genres_urls.values()).index(url)]
    write_summary_to_doc(summary_text, genre)
